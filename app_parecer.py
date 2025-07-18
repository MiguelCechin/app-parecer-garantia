import streamlit as st
import docx
from docx import Document
from docx.shared import Pt
from io import BytesIO

def avaliar_titulo(respostas):
    r11 = respostas.get("1.1", "").upper()
    if r11 == "NÃO":
        return "  2.  Não foi apresentado, até o momento, instrumento formal de constituição da dívida, impossibilitando a verificação de sua regularidade e eventual cessibilidade."
    elif r11 == "NÃO SE APLICA":
        return "  2.  A análise do título/instrumento de formalização do crédito não se aplica à presente garantia."
    elif r11 == "SIM":
        r12 = respostas.get("1.2", "").upper()
        if r12 == "NÃO":
            return "  2.  Inicialmente, verificou-se a existência de documento representativo de dívida que carece de assinatura e não preenche os requisitos formais mínimos aplicáveis."
        elif r12 == "SIM":
            r13 = respostas.get("1.3", "").upper()
            if r13 == "SIM":
                return "  2.  Inicialmente, verificou-se a existência de documento representativo de dívida que carece da assinatura de testemunhas, e portanto não preenche os requisitos formais mínimos aplicáveis."
            elif r13 == "NÃO":
                r14 = respostas.get("1.4", "").upper()
                if r14 == "NÃO":
                    return "  2.  Inicialmente, verificou-se a existência de documento representativo de dívida que carece da assinatura de testemunhas, e portanto não preenche os requisitos formais mínimos aplicáveis."
                elif r14 == "SIM":
                    return "  2.  Inicialmente, verificou-se a existência de documento representativo de dívida com o preenchimento de requisitos formais mínimos aplicáveis."
    return "Resposta inválida ou incompleta."

def gerar_parecer_garantia(dados):
    doc = Document()
    # Configurar fonte padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = 1
    titulo.add_run("PARECER SIMPLIFICADO").bold = True
    doc.add_paragraph()

    # Subtítulo
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = 1
    subtitulo.add_run("Para fins de monitoramento de garantias")
    doc.add_paragraph()

    # Parágrafo inicial
    para1 = doc.add_paragraph()
    para1.alignment = 3
    para1.add_run(
        f"Em atenção à solicitação feita pela {dados['solicitante']} na qualidade de gestora do(s) {dados['gestora']}, apresentamos o presente parecer jurídico simplificado a respeito da capacidade de execução de garantias ligadas a ativos financeiros representativos de dívidas ou obrigações titularizados pelo(s) Fundo(s)."
    )
    doc.add_paragraph()

    # Seção A
    doc.add_paragraph().add_run("A) INFORMAÇÕES PRELIMINARES").bold = True
    doc.add_paragraph().add_run("1. Para a elaboração deste parecer, foram acessadas as seguintes informações e/ou documentos:")
    doc.add_paragraph()

    table1 = doc.add_table(rows=3, cols=2)
    table1.style = 'Table Grid'
    table1.cell(0,0).text = "TÍTULO"
    table1.cell(0,1).text = f"CCB nº {dados['numero_celula']}"
    table1.cell(1,0).text = "GARANTIA"
    table1.cell(1,1).text = "Alienação fiduciária de veículo"
    table1.cell(2,0).text = "DOCUMENTOS RECEBIDOS"
    table1.cell(2,1).text = "\n".join(dados['docs'])
    doc.add_paragraph()

    # Seção B
    doc.add_paragraph().add_run("B) DADOS BÁSICOS DA OPERAÇÃO").bold = True
    table2 = doc.add_table(rows=7, cols=2)
    table2.style = 'Table Grid'
    table2.cell(0,0).text = "CESSIONÁRIO"
    table2.cell(0,1).text = dados['cessionario']
    table2.cell(1,0).text = "CEDENTE"
    table2.cell(1,1).text = dados['cedente']
    table2.cell(2,0).text = "DEVEDOR(A)"
    table2.cell(2,1).text = dados['devedor']
    table2.cell(3,0).text = "DATA DE EMISSÃO/REFERÊNCIA"
    table2.cell(3,1).text = dados['data_emissao']
    table2.cell(4,0).text = "VALOR DA OPERAÇÃO"
    table2.cell(4,1).text = f"R$ {dados['valor_operacao']}"
    table2.cell(5,0).text = "DATA DA PARCELA 1"
    table2.cell(5,1).text = dados['data_primeira_parcela']
    table2.cell(6,0).text = "DATA DA PARCELA FINAL"
    table2.cell(6,1).text = dados['data_ultima_parcela']
    doc.add_paragraph()

    # Seção C
    doc.add_paragraph().add_run("C) GARANTIA").bold = True
    table3 = doc.add_table(rows=5, cols=2)
    table3.style = 'Table Grid'
    table3.cell(0,0).text = "FIDUCIANTE"
    table3.cell(0,1).text = dados['fiduciante']
    table3.cell(1,0).text = "FIDUCIÁRIO(A) ORIGINAL"
    table3.cell(1,1).text = dados['fiduciario']
    table3.cell(2,0).text = "BEM OBJETO DE GARANTIA"
    table3.cell(2,1).text = dados['obj_garantia']
    table3.cell(3,0).text = "VALOR DE AVALIAÇÃO HISTÓRICO"
    table3.cell(3,1).text = f"R$ {dados['valor_obj_inicial']}"
    table3.cell(4,0).text = "VALOR DE AVALIAÇÃO ATUAL"
    table3.cell(4,1).text = f"R$ {dados['valor_obj_atual']}"
    doc.add_paragraph()

    # Seção D
    doc.add_paragraph().add_run("D) PRINCIPAIS CONSTATAÇÕES E APONTAMENTOS").bold = True
    resultado = avaliar_titulo(dados['respostas_titulo'])
    doc.add_paragraph().add_run(resultado)

    return doc


def main():
    st.title("Gerador de Parecer Jurídico Simplificado")

    st.header("Dados da Operação")
    solicitante = st.text_input("Solicitante")
    gestora = st.text_input("Gestora")
    numero_celula = st.text_input("Número da Célula (CCB)")
    docs_input = st.text_area("Documentos Recebidos (um por linha)")
    cessionario = st.text_input("Cessionário")
    cedente = st.text_input("Cedente")
    devedor = st.text_input("Devedor(a)")
    data_emissao = st.date_input("Data de Emissão/Referência")
    valor_operacao = st.text_input("Valor da Operação")
    data_primeira = st.date_input("Data da Parcela 1")
    data_ultima = st.date_input("Data da Parcela Final")

    st.header("Dados da Garantia")
    fiduciante = st.text_input("Fiduciante")
    fiduciario = st.text_input("Fiduciário(a) Original")
    obj_garantia = st.text_input("Informações do Objeto de Garantia")
    valor_inicial = st.text_input("Valor de Avaliação Histórico")
    valor_atual = st.text_input("Valor de Avaliação Atual")

    st.header("Título")
    r11 = st.radio("1.1 - Foi submetido à análise título/instrumento de formalização do crédito?", ["SIM", "NÃO", "NÃO SE APLICA"])
    r12 = st.radio("1.2 - O documento mencionado neste item está assinado?", ["SIM", "NÃO"])
    r13 = st.radio("1.3 - A assinatura é eletrônica?", ["SIM", "NÃO"])
    r14 = st.radio("1.4 - Em não sendo eletrônica, o documento está assinado por 2 testemunhas?", ["SIM", "NÃO"])

    if st.button("Gerar e Baixar Parecer"):
        dados = {
            'solicitante': solicitante,
            'gestora': gestora,
            'numero_celula': numero_celula,
            'docs': docs_input.splitlines(),
            'cessionario': cessionario,
            'cedente': cedente,
            'devedor': devedor,
            'data_emissao': data_emissao.strftime("%d/%m/%Y"),
            'valor_operacao': valor_operacao,
            'data_primeira_parcela': data_primeira.strftime("%d/%m/%Y"),
            'data_ultima_parcela': data_ultima.strftime("%d/%m/%Y"),
            'fiduciante': fiduciante,
            'fiduciario': fiduciario,
            'obj_garantia': obj_garantia,
            'valor_obj_inicial': valor_inicial,
            'valor_obj_atual': valor_atual,
            'respostas_titulo': {'1.1': r11, '1.2': r12, '1.3': r13, '1.4': r14}
        }
        doc = gerar_parecer_garantia(dados)
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        st.download_button(
            label="Baixar Parecer (DOCX)",
            data=output,
            file_name=f"Parecer_CCB_{numero_celula}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()


